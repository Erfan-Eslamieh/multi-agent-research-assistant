from fpdf import FPDF
from docx import Document
import io

def export_to_pdf(content: str, title: str = "Export") -> bytes:
    pdf = FPDF()
    pdf.set_margins(20, 20, 20)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, title, ln=True)
    pdf.ln(4)
    pdf.set_font("Helvetica", size=9)
    
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            pdf.ln(3)
            continue
        # فقط ASCII نگه دار
        line = line.encode('ascii', 'ignore').decode('ascii')
        if not line.strip():
            pdf.ln(3)
            continue
        if len(line) > 200:
            line = line[:200] + "..."
        pdf.multi_cell(0, 5, line)
    
    return pdf.output()

def export_to_word(content: str, title: str = "Export") -> bytes:
    doc = Document()
    doc.add_heading(title, 0)
    
    for line in content.split("\n"):
        line = line.strip()
        if line.startswith("**") and line.endswith("**"):
            doc.add_heading(line.replace("**", ""), level=2)
        elif line:
            doc.add_paragraph(line)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()

def chat_to_text(chat_history: list) -> str:
    lines = []
    for msg in chat_history:
        if msg["role"] == "user":
            lines.append(f"Question: {msg['content']}\n")
        else:
            lines.append(f"Answer: {msg['content']}\n")
            if "details" in msg:
                lines.append(f"Confidence: {msg['details'].get('confidence', '')}%")
                lines.append(f"Source pages: {', '.join(msg['details'].get('source_pages', []))}\n")
    return "\n".join(lines)